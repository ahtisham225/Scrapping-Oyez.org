import docx
import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
DRIVER_PATH = 'chromedriver'
# driver = webdriver.Chrome(executable_path=DRIVER_PATH)
# driver.get('https://google.com')
# from selenium import webdriver
u = 19300
while u < 19400:

    options = Options()
    options.headless = True
    options.add_argument("--window-size=1920,1200")

    driver = webdriver.Chrome(options=options, executable_path=DRIVER_PATH)
    driver.get(
        "https://apps.oyez.org/player/#/burger8/oral_argument_audio/" + str(u))
    # print(driver.page_source)
    # print("https://apps.oyez.org/player/#/burger8/oral_argument_audio/" + str(u))
    soup = BeautifulSoup(driver.page_source, 'html.parser')
    # print(soup)
    # print(driver)

    doc = docx.Document()
    divisions = soup.find_all(class_="ng-binding")
    names = soup.find_all('h4')
    c = set()
    for i in names:
        c.add(i.string)
    # print(c)

    # print(divisions.get_text())
    j = 0
    title = ""
    for i in divisions:
        # print(i.string)
        if j == 0:
            title = i.string
            doc.add_heading(i.string, 0)
            j = j + 1
        elif j == 1 or j == 2:
            j = j+1
            pass
        elif i.string != None:
            if i.string in c:
                # print('a')
                doc.add_heading(i.string, level=4)
            else:
                doc.add_paragraph(i.string)
    # print(u)
    print(title)
    print(u)
    if title == 'Oyez media player':
        pass
    else:
        doc.save(title + '.docx')
        u += 1
    driver.quit()
# names = soup.find_all('h4')
# print(names)
# for j in divisions:
#     print(j.get_text())
